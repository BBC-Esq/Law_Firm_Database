import os
import sys
import calendar
from datetime import date
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QComboBox,
    QDoubleSpinBox, QPushButton, QGroupBox, QFormLayout, QMessageBox, QFileDialog
)
from PySide6.QtCore import Qt
from core.utils import format_matter_display
from gui.utils import select_all_on_focus, load_combo_with_items
from gui.widgets.styled_combo_box import StyledComboBox

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def get_image_path():
    if getattr(sys, 'frozen', False):
        return os.path.join(os.path.dirname(sys.executable), "law_image.jpg")
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", "law_image.jpg")


def set_cell_border(cell, border_size=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_no_paragraph_spacing(doc):
    style = doc.styles['Normal']
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)


def add_paragraph_no_spacing(doc, text="", alignment=None, bold=False, font_size=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    if alignment:
        para.alignment = alignment
    if text:
        run = para.add_run(text)
        run.bold = bold
        if font_size:
            run.font.size = font_size
    return para


class InvoiceWidget(QWidget):
    def __init__(self, case_queries, billing_queries, invoice_queries, get_show_closed_callback=None):
        super().__init__()
        self.case_queries = case_queries
        self.billing_queries = billing_queries
        self.invoice_queries = invoice_queries
        self.get_show_closed = get_show_closed_callback or (lambda: True)
        self.image_path = get_image_path()
        self.matters = []
        self.setup_ui()
        self.load_matters()

    def setup_ui(self):
        layout = QVBoxLayout(self)

        if not HAS_DOCX:
            warning_label = QLabel("python-docx is required for invoice generation.\nInstall with: pip install python-docx")
            warning_label.setStyleSheet("color: red; font-weight: bold; padding: 20px;")
            layout.addWidget(warning_label)

        matter_group = QGroupBox("Select Matter")
        matter_layout = QVBoxLayout(matter_group)
        self.matter_combo = StyledComboBox()
        self.matter_combo.setMinimumWidth(400)
        matter_layout.addWidget(self.matter_combo)
        layout.addWidget(matter_group)

        period_group = QGroupBox("Invoice Period")
        period_layout = QHBoxLayout(period_group)
        period_layout.addWidget(QLabel("Month:"))
        self.month_combo = QComboBox()
        for i in range(1, 13):
            self.month_combo.addItem(date(2000, i, 1).strftime("%B"), i)
        self.month_combo.setCurrentIndex(date.today().month - 1)
        period_layout.addWidget(self.month_combo)
        period_layout.addSpacing(20)
        period_layout.addWidget(QLabel("Year:"))
        self.year_combo = QComboBox()
        current_year = date.today().year
        for year in range(current_year - 5, current_year + 2):
            self.year_combo.addItem(str(year), year)
        self.year_combo.setCurrentText(str(current_year))
        period_layout.addWidget(self.year_combo)
        period_layout.addStretch()
        layout.addWidget(period_group)

        trust_group = QGroupBox("Trust Account Replenishment Targets")
        trust_layout = QFormLayout(trust_group)
        self.fee_target_spin = QDoubleSpinBox()
        self.fee_target_spin.setRange(0, 1000000)
        self.fee_target_spin.setDecimals(2)
        self.fee_target_spin.setPrefix("$")
        self.fee_target_spin.setValue(0)
        select_all_on_focus(self.fee_target_spin)
        trust_layout.addRow("Fee Trust Target:", self.fee_target_spin)
        self.expense_target_spin = QDoubleSpinBox()
        self.expense_target_spin.setRange(0, 1000000)
        self.expense_target_spin.setDecimals(2)
        self.expense_target_spin.setPrefix("$")
        self.expense_target_spin.setValue(0)
        select_all_on_focus(self.expense_target_spin)
        trust_layout.addRow("Expense Trust Target:", self.expense_target_spin)
        layout.addWidget(trust_group)

        reconcile_group = QGroupBox("Trust Account Reconciliation")
        reconcile_layout = QVBoxLayout(reconcile_group)

        self.reconcile_combo = QComboBox()
        self.reconcile_combo.addItem("Keep Separate (No Cross-Account Transfer)", "none")
        self.reconcile_combo.addItem("Final Invoice (Combine All Balances)", "final")
        self.reconcile_combo.addItem("Transfer Surplus & Replenish", "transfer")
        self.reconcile_combo.currentIndexChanged.connect(self.on_reconcile_mode_changed)
        reconcile_layout.addWidget(self.reconcile_combo)

        self.reconcile_description = QLabel()
        self.reconcile_description.setWordWrap(True)
        self.reconcile_description.setStyleSheet("color: #666; font-size: 9pt; padding: 5px;")
        reconcile_layout.addWidget(self.reconcile_description)
        self.on_reconcile_mode_changed()

        layout.addWidget(reconcile_group)

        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        self.generate_btn = QPushButton("Generate Invoice")
        self.generate_btn.setMinimumWidth(200)
        self.generate_btn.clicked.connect(self.generate_invoice)
        self.generate_btn.setEnabled(HAS_DOCX)
        btn_layout.addWidget(self.generate_btn)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        layout.addStretch()

    def on_reconcile_mode_changed(self):
        mode = self.reconcile_combo.currentData()
        descriptions = {
            "none": "Fee and expense accounts are tracked separately. Each account's replenishment is calculated independently.",
            "final": "For closing out a client's account. Combines fee and expense balances to calculate a single final amount due. Any surplus in one account offsets deficit in the other.",
            "transfer": "Applies any surplus from the expense account to cover fee deficits (or vice versa), then requires replenishment of both accounts back to their target levels."
        }
        self.reconcile_description.setText(descriptions.get(mode, ""))

    def load_matters(self):
        include_closed = self.get_show_closed()
        self.matters = self.case_queries.get_matters_for_invoice(include_closed=include_closed)
        load_combo_with_items(
            self.matter_combo,
            self.matters,
            lambda m: (format_matter_display(m, include_client=True), m),
            "-- Select a Matter --"
        )

    def refresh(self):
        self.load_matters()

    def set_column_widths(self, table, widths):
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = width

    def remove_paragraph_spacing_in_table(self, table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)

    def add_header(self, doc):
        section = doc.sections[0]
        section.different_first_page_header_footer = True
        header = section.first_page_header
        header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        header_table.autofit = False
        left_cell = header_table.rows[0].cells[0]
        right_cell = header_table.rows[0].cells[1]
        set_cell_border(left_cell)
        set_cell_border(right_cell)
        left_cell.width = Inches(3.25)
        right_cell.width = Inches(3.25)
        if os.path.exists(self.image_path):
            left_para = left_cell.paragraphs[0]
            left_para.paragraph_format.space_before = Pt(0)
            left_para.paragraph_format.space_after = Pt(0)
            left_run = left_para.add_run()
            left_run.add_picture(self.image_path, width=Inches(2.5))
        right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        right_para.paragraph_format.space_before = Pt(0)
        right_para.paragraph_format.space_after = Pt(0)
        phone_run = right_para.add_run("404-556-7057")
        phone_run.font.size = Pt(10)
        right_para.add_run("\n")
        email_run = right_para.add_run("bbc@chintellalaw.com")
        email_run.font.size = Pt(10)
        right_para.add_run("\n")
        web_run = right_para.add_run("www.chintellalaw.com")
        web_run.font.size = Pt(10)

    def calculate_reconciliation(self, fee_balance, expense_balance, fee_target, expense_target, mode):
        result = {
            'original_fee_balance': fee_balance,
            'original_expense_balance': expense_balance,
            'transfer_amount': 0,
            'transfer_direction': None,
            'is_final': False
        }

        if mode == "final":
            result['is_final'] = True
            combined_balance = fee_balance + expense_balance
            result['total_due'] = max(0, -combined_balance)
            result['fee_replenishment'] = 0
            result['expense_replenishment'] = 0
            result['adjusted_fee_balance'] = fee_balance
            result['adjusted_expense_balance'] = expense_balance
            if expense_balance > 0 and fee_balance < 0:
                result['transfer_amount'] = min(expense_balance, -fee_balance)
                result['transfer_direction'] = 'expense_to_fee'
            elif fee_balance > 0 and expense_balance < 0:
                result['transfer_amount'] = min(fee_balance, -expense_balance)
                result['transfer_direction'] = 'fee_to_expense'

        elif mode == "transfer":
            adjusted_fee_balance = fee_balance
            adjusted_expense_balance = expense_balance

            if expense_balance > 0 and fee_balance < 0:
                transfer = min(expense_balance, -fee_balance)
                result['transfer_amount'] = transfer
                result['transfer_direction'] = 'expense_to_fee'
                adjusted_fee_balance = fee_balance + transfer
                adjusted_expense_balance = expense_balance - transfer
            elif fee_balance > 0 and expense_balance < 0:
                transfer = min(fee_balance, -expense_balance)
                result['transfer_amount'] = transfer
                result['transfer_direction'] = 'fee_to_expense'
                adjusted_fee_balance = fee_balance - transfer
                adjusted_expense_balance = expense_balance + transfer

            result['adjusted_fee_balance'] = adjusted_fee_balance
            result['adjusted_expense_balance'] = adjusted_expense_balance
            result['fee_replenishment'] = max(0, fee_target - adjusted_fee_balance)
            result['expense_replenishment'] = max(0, expense_target - adjusted_expense_balance)
            result['total_due'] = result['fee_replenishment'] + result['expense_replenishment']

        else:
            result['adjusted_fee_balance'] = fee_balance
            result['adjusted_expense_balance'] = expense_balance
            result['fee_replenishment'] = max(0, fee_target - fee_balance)
            result['expense_replenishment'] = max(0, expense_target - expense_balance)
            result['total_due'] = result['fee_replenishment'] + result['expense_replenishment']

        return result

    def generate_invoice(self):
        if not HAS_DOCX:
            QMessageBox.warning(self, "Missing Dependency", "python-docx is required for invoice generation.\nInstall with: pip install python-docx")
            return

        matter = self.matter_combo.currentData()
        if not matter:
            QMessageBox.warning(self, "Warning", "Please select a matter.")
            return

        case_id = matter['id']
        year = self.year_combo.currentData()
        month = self.month_combo.currentData()
        fee_target = self.fee_target_spin.value()
        expense_target = self.expense_target_spin.value()
        reconcile_mode = self.reconcile_combo.currentData()

        last_day = calendar.monthrange(year, month)[1]
        balance_date_str = f"{month}/{last_day}/{year}"

        entries = self.billing_queries.get_entries_for_period(case_id, year, month)
        billing_rate = self.invoice_queries.get_billing_rate(case_id)
        trust_data = self.invoice_queries.get_trust_balances(case_id, year, month)

        fee_balance = trust_data['fee_balance']
        expense_balance = trust_data['expense_balance']

        reconciliation = self.calculate_reconciliation(
            fee_balance, expense_balance, fee_target, expense_target, reconcile_mode
        )

        time_entries = [e for e in entries if not e['is_expense']]
        expense_entries = [e for e in entries if e['is_expense']]
        period_fees = sum((e['hours'] or 0) * billing_rate for e in time_entries)
        period_expenses = sum((e['amount_cents'] or 0) / 100.0 for e in expense_entries)

        client_name = f"{matter.get('first_name') or ''} {matter.get('last_name') or ''}".strip() or "Client"
        client_address = matter.get('address') or ''
        case_name = matter.get('case_name') or ''
        period_name = date(year, month, 1).strftime("%B %Y")

        doc = Document()
        set_no_paragraph_spacing(doc)
        self.add_header(doc)
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)

        invoice_type = "FINAL INVOICE" if reconciliation['is_final'] else "INVOICE"
        add_paragraph_no_spacing(doc, invoice_type, WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=Pt(18))
        add_paragraph_no_spacing(doc)

        info_table = doc.add_table(rows=1, cols=2)
        info_table.autofit = True
        left_cell = info_table.rows[0].cells[0]
        left_cell.text = f"Bill To:\n{client_name}"
        if client_address:
            left_cell.text += f"\n{client_address}"
        right_cell = info_table.rows[0].cells[1]
        right_cell.text = f"Invoice Date: {date.today().strftime('%B %d, %Y')}\nPeriod: {period_name}\nMatter: {case_name}"
        self.remove_paragraph_spacing_in_table(info_table)
        add_paragraph_no_spacing(doc)

        if time_entries:
            add_paragraph_no_spacing(doc, "Professional Services", bold=True, font_size=Pt(12))
            fees_table = doc.add_table(rows=1, cols=5)
            fees_table.style = 'Table Grid'
            fees_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            fees_table.autofit = False
            self.set_column_widths(fees_table, [Inches(0.9), Inches(3.6), Inches(0.6), Inches(0.7), Inches(0.8)])
            hdr_cells = fees_table.rows[0].cells
            headers = ['Date', 'Description', 'Hours', 'Rate', 'Amount']
            for i, text in enumerate(headers):
                hdr_cells[i].text = text
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for entry in time_entries:
                row_cells = fees_table.add_row().cells
                row_cells[0].text = str(entry['entry_date'])
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row_cells[1].text = entry['description'] or ''
                row_cells[2].text = f"{entry['hours']:.1f}"
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row_cells[3].text = f"${billing_rate:.2f}"
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[3].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                amount = (entry['hours'] or 0) * billing_rate
                row_cells[4].text = f"${amount:.2f}"
                row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[4].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            total_row = fees_table.add_row().cells
            total_row[0].merge(total_row[3])
            total_row[0].text = "Total Professional Services"
            total_row[0].paragraphs[0].runs[0].bold = True
            total_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            total_row[4].text = f"${period_fees:.2f}"
            total_row[4].paragraphs[0].runs[0].bold = True
            total_row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.set_column_widths(fees_table, [Inches(0.9), Inches(3.6), Inches(0.6), Inches(0.7), Inches(0.8)])
            self.remove_paragraph_spacing_in_table(fees_table)
            add_paragraph_no_spacing(doc)

        if expense_entries:
            add_paragraph_no_spacing(doc, "Expenses", bold=True, font_size=Pt(12))
            exp_table = doc.add_table(rows=1, cols=3)
            exp_table.style = 'Table Grid'
            exp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            exp_table.autofit = False
            self.set_column_widths(exp_table, [Inches(0.9), Inches(4.9), Inches(0.8)])
            hdr_cells = exp_table.rows[0].cells
            headers = ['Date', 'Description', 'Amount']
            for i, text in enumerate(headers):
                hdr_cells[i].text = text
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for entry in expense_entries:
                row_cells = exp_table.add_row().cells
                row_cells[0].text = str(entry['entry_date'])
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row_cells[1].text = entry['description'] or ''
                amount = (entry['amount_cents'] or 0) / 100.0
                row_cells[2].text = f"${amount:.2f}"
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            total_row = exp_table.add_row().cells
            total_row[0].merge(total_row[1])
            total_row[0].text = "Total Expenses"
            total_row[0].paragraphs[0].runs[0].bold = True
            total_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            total_row[2].text = f"${period_expenses:.2f}"
            total_row[2].paragraphs[0].runs[0].bold = True
            total_row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.set_column_widths(exp_table, [Inches(0.9), Inches(4.9), Inches(0.8)])
            self.remove_paragraph_spacing_in_table(exp_table)
            add_paragraph_no_spacing(doc)

        add_paragraph_no_spacing(doc, f"Trust Account Summary as of {balance_date_str}", bold=True, font_size=Pt(12))

        if reconciliation['is_final']:
            rows_data = [
                ("FEE TRUST ACCOUNT", ""),
                ("Total Fee Payments Received:", f"${trust_data['total_fee_payments']:.2f}"),
                ("Total Fees Billed:", f"(${trust_data['total_fees_billed']:.2f})"),
                ("Fee Trust Balance:", f"${fee_balance:.2f}"),
                ("", ""),
                ("EXPENSE TRUST ACCOUNT", ""),
                ("Total Expense Payments Received:", f"${trust_data['total_expense_payments']:.2f}"),
                ("Total Expenses Billed:", f"(${trust_data['total_expenses_billed']:.2f})"),
                ("Expense Trust Balance:", f"${expense_balance:.2f}"),
                ("", ""),
                ("Combined Trust Balance:", f"${fee_balance + expense_balance:.2f}"),
            ]

            if reconciliation['transfer_amount'] > 0:
                if reconciliation['transfer_direction'] == 'expense_to_fee':
                    rows_data.append((f"(Expense funds applied to fees: ${reconciliation['transfer_amount']:.2f})", ""))
                else:
                    rows_data.append((f"(Fee funds applied to expenses: ${reconciliation['transfer_amount']:.2f})", ""))

            summary_table = doc.add_table(rows=len(rows_data), cols=2)
            summary_table.style = 'Table Grid'
            summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, (label, value) in enumerate(rows_data):
                summary_table.rows[i].cells[0].text = label
                summary_table.rows[i].cells[1].text = value
                summary_table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if label in ("FEE TRUST ACCOUNT", "EXPENSE TRUST ACCOUNT", "Fee Trust Balance:", "Expense Trust Balance:", "Combined Trust Balance:"):
                    for paragraph in summary_table.rows[i].cells[0].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                    if value:
                        for paragraph in summary_table.rows[i].cells[1].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
            self.remove_paragraph_spacing_in_table(summary_table)

        elif reconciliation['transfer_amount'] > 0:
            if reconciliation['transfer_direction'] == 'expense_to_fee':
                transfer_label = "Transfer from Expense to Fee Trust:"
            else:
                transfer_label = "Transfer from Fee to Expense Trust:"

            rows_data = [
                ("FEE TRUST ACCOUNT", ""),
                ("Total Fee Payments Received:", f"${trust_data['total_fee_payments']:.2f}"),
                ("Total Fees Billed:", f"(${trust_data['total_fees_billed']:.2f})"),
                ("Fee Trust Balance:", f"${fee_balance:.2f}"),
                ("", ""),
                ("EXPENSE TRUST ACCOUNT", ""),
                ("Total Expense Payments Received:", f"${trust_data['total_expense_payments']:.2f}"),
                ("Total Expenses Billed:", f"(${trust_data['total_expenses_billed']:.2f})"),
                ("Expense Trust Balance:", f"${expense_balance:.2f}"),
                ("", ""),
                (transfer_label, f"${reconciliation['transfer_amount']:.2f}"),
                ("", ""),
                ("Adjusted Fee Trust Balance:", f"${reconciliation['adjusted_fee_balance']:.2f}"),
                ("Fee Trust Target:", f"${fee_target:.2f}"),
                ("Fee Replenishment Required:", f"${reconciliation['fee_replenishment']:.2f}"),
                ("", ""),
                ("Adjusted Expense Trust Balance:", f"${reconciliation['adjusted_expense_balance']:.2f}"),
                ("Expense Trust Target:", f"${expense_target:.2f}"),
                ("Expense Replenishment Required:", f"${reconciliation['expense_replenishment']:.2f}"),
            ]

            summary_table = doc.add_table(rows=len(rows_data), cols=2)
            summary_table.style = 'Table Grid'
            summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, (label, value) in enumerate(rows_data):
                summary_table.rows[i].cells[0].text = label
                summary_table.rows[i].cells[1].text = value
                summary_table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if label in ("FEE TRUST ACCOUNT", "EXPENSE TRUST ACCOUNT", "Fee Trust Balance:", "Expense Trust Balance:", "Fee Replenishment Required:", "Expense Replenishment Required:"):
                    for paragraph in summary_table.rows[i].cells[0].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                    if value:
                        for paragraph in summary_table.rows[i].cells[1].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
            self.remove_paragraph_spacing_in_table(summary_table)
        else:
            rows_data = [
                ("FEE TRUST ACCOUNT", ""),
                ("Total Fee Payments Received:", f"${trust_data['total_fee_payments']:.2f}"),
                ("Total Fees Billed:", f"(${trust_data['total_fees_billed']:.2f})"),
                ("Fee Trust Balance:", f"${fee_balance:.2f}"),
                ("Fee Trust Target:", f"${fee_target:.2f}"),
                ("Fee Replenishment Required:", f"${reconciliation['fee_replenishment']:.2f}"),
                ("", ""),
                ("EXPENSE TRUST ACCOUNT", ""),
                ("Total Expense Payments Received:", f"${trust_data['total_expense_payments']:.2f}"),
                ("Total Expenses Billed:", f"(${trust_data['total_expenses_billed']:.2f})"),
                ("Expense Trust Balance:", f"${expense_balance:.2f}"),
                ("Expense Trust Target:", f"${expense_target:.2f}"),
                ("Expense Replenishment Required:", f"${reconciliation['expense_replenishment']:.2f}"),
            ]

            summary_table = doc.add_table(rows=len(rows_data), cols=2)
            summary_table.style = 'Table Grid'
            summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, (label, value) in enumerate(rows_data):
                summary_table.rows[i].cells[0].text = label
                summary_table.rows[i].cells[1].text = value
                summary_table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if label in ("FEE TRUST ACCOUNT", "EXPENSE TRUST ACCOUNT", "Fee Trust Balance:", "Expense Trust Balance:", "Fee Replenishment Required:", "Expense Replenishment Required:"):
                    for paragraph in summary_table.rows[i].cells[0].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                    if value:
                        for paragraph in summary_table.rows[i].cells[1].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
            self.remove_paragraph_spacing_in_table(summary_table)

        add_paragraph_no_spacing(doc)
        add_paragraph_no_spacing(doc, f"TOTAL AMOUNT DUE: ${reconciliation['total_due']:.2f}", WD_ALIGN_PARAGRAPH.RIGHT, bold=True, font_size=Pt(14))
        add_paragraph_no_spacing(doc)
        add_paragraph_no_spacing(doc)
        add_paragraph_no_spacing(doc, "Thank you for your business.", WD_ALIGN_PARAGRAPH.CENTER)

        default_filename = f"Invoice_{case_name}_{period_name.replace(' ', '_')}.docx"
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save Invoice", default_filename, "Word Documents (*.docx)"
        )
        if file_path:
            doc.save(file_path)
            QMessageBox.information(self, "Success", f"Invoice saved to:\n{file_path}")