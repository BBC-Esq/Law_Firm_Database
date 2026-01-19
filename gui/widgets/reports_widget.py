import os
import sys
from datetime import date
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QComboBox,
    QPushButton, QGroupBox, QTableWidget, QTableWidgetItem,
    QHeaderView, QMessageBox, QFileDialog, QCheckBox
)
from PySide6.QtCore import Qt
from PySide6.QtGui import QColor, QBrush

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def get_image_path():
    if getattr(sys, 'frozen', False):
        return os.path.join(os.path.dirname(sys.executable), "law_image.jpg")
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", "law_image.jpg")


class ReportsWidget(QWidget):
    HEADERS = [
        "Matter", "Client", "Status", "Hours", "Fees Billed", 
        "Expenses Billed", "Total Billed", "Fee Payments", 
        "Expense Payments", "Total Payments", "Balance"
    ]

    def __init__(self, report_queries, get_show_closed_callback=None):
        super().__init__()
        self.report_queries = report_queries
        self.get_show_closed = get_show_closed_callback or (lambda: True)
        self.image_path = get_image_path()
        self.current_data = []
        self.setup_ui()

    def setup_ui(self):
        layout = QVBoxLayout(self)

        controls_group = QGroupBox("Report Settings")
        controls_layout = QHBoxLayout(controls_group)

        controls_layout.addWidget(QLabel("Report Type:"))
        self.report_type_combo = QComboBox()
        self.report_type_combo.addItem("Monthly Billing Summary", "monthly")
        self.report_type_combo.addItem("All-Time Summary", "all_time")
        self.report_type_combo.currentIndexChanged.connect(self.on_report_type_changed)
        controls_layout.addWidget(self.report_type_combo)

        controls_layout.addSpacing(20)

        controls_layout.addWidget(QLabel("Month:"))
        self.month_combo = QComboBox()
        for i in range(1, 13):
            self.month_combo.addItem(date(2000, i, 1).strftime("%B"), i)
        self.month_combo.setCurrentIndex(date.today().month - 1)
        controls_layout.addWidget(self.month_combo)

        controls_layout.addWidget(QLabel("Year:"))
        self.year_combo = QComboBox()
        current_year = date.today().year
        for year in range(current_year - 5, current_year + 2):
            self.year_combo.addItem(str(year), year)
        self.year_combo.setCurrentText(str(current_year))
        controls_layout.addWidget(self.year_combo)

        controls_layout.addSpacing(20)

        self.include_closed_checkbox = QCheckBox("Include Closed Matters")
        self.include_closed_checkbox.setChecked(True)
        controls_layout.addWidget(self.include_closed_checkbox)

        controls_layout.addStretch()

        self.generate_btn = QPushButton("Generate Report")
        self.generate_btn.clicked.connect(self.generate_report)
        controls_layout.addWidget(self.generate_btn)

        layout.addWidget(controls_group)

        totals_group = QGroupBox("Summary")
        totals_layout = QHBoxLayout(totals_group)

        self.matters_label = QLabel("Matters: --")
        self.matters_label.setStyleSheet("font-weight: bold;")
        totals_layout.addWidget(self.matters_label)

        totals_layout.addSpacing(20)

        self.hours_label = QLabel("Hours: --")
        totals_layout.addWidget(self.hours_label)

        totals_layout.addSpacing(20)

        self.billed_label = QLabel("Total Billed: --")
        totals_layout.addWidget(self.billed_label)

        totals_layout.addSpacing(20)

        self.payments_label = QLabel("Total Payments: --")
        totals_layout.addWidget(self.payments_label)

        totals_layout.addSpacing(20)

        self.balance_label = QLabel("Net Balance: --")
        self.balance_label.setStyleSheet("font-weight: bold;")
        totals_layout.addWidget(self.balance_label)

        totals_layout.addStretch()
        layout.addWidget(totals_group)

        self.table = QTableWidget()
        self.table.setColumnCount(len(self.HEADERS))
        self.table.setHorizontalHeaderLabels(self.HEADERS)
        self.table.setSelectionBehavior(QTableWidget.SelectRows)
        self.table.setSelectionMode(QTableWidget.SingleSelection)
        self.table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.table.setAlternatingRowColors(True)
        self.table.setSortingEnabled(True)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        layout.addWidget(self.table)

        export_layout = QHBoxLayout()
        export_layout.addStretch()

        self.export_docx_btn = QPushButton("Export to Word")
        self.export_docx_btn.clicked.connect(self.export_to_docx)
        self.export_docx_btn.setEnabled(HAS_DOCX)
        export_layout.addWidget(self.export_docx_btn)

        self.export_csv_btn = QPushButton("Export to CSV")
        self.export_csv_btn.clicked.connect(self.export_to_csv)
        export_layout.addWidget(self.export_csv_btn)

        export_layout.addStretch()
        layout.addLayout(export_layout)

    def on_report_type_changed(self, index):
        is_monthly = self.report_type_combo.currentData() == "monthly"
        self.month_combo.setEnabled(is_monthly)
        self.year_combo.setEnabled(is_monthly)

    def generate_report(self):
        report_type = self.report_type_combo.currentData()
        include_closed = self.include_closed_checkbox.isChecked()

        if report_type == "monthly":
            year = self.year_combo.currentData()
            month = self.month_combo.currentData()
            self.current_data = self.report_queries.get_monthly_billing_summary(year, month, include_closed)
            totals = self.report_queries.get_period_totals(year, month, include_closed)
        else:
            self.current_data = self.report_queries.get_all_matters_summary(include_closed)
            totals = self._calculate_totals(self.current_data)

        self.populate_table(self.current_data)
        self.update_totals(totals)

    def _calculate_totals(self, data: list) -> dict:
        totals = {
            'total_hours': 0,
            'total_fees_cents': 0,
            'total_expenses_cents': 0,
            'total_fee_payments_cents': 0,
            'total_expense_payments_cents': 0,
            'matter_count': len(data)
        }

        for row in data:
            totals['total_hours'] += row.get('total_hours') or 0
            totals['total_fees_cents'] += row.get('total_fees_cents') or 0
            totals['total_expenses_cents'] += row.get('total_expenses_cents') or 0
            totals['total_fee_payments_cents'] += row.get('total_fee_payments_cents') or 0
            totals['total_expense_payments_cents'] += row.get('total_expense_payments_cents') or 0

        totals['total_billed_cents'] = totals['total_fees_cents'] + totals['total_expenses_cents']
        totals['total_payments_cents'] = totals['total_fee_payments_cents'] + totals['total_expense_payments_cents']

        return totals

    def populate_table(self, data: list):
        self.table.setSortingEnabled(False)
        self.table.setRowCount(len(data))

        for row_idx, row in enumerate(data):
            total_hours = row.get('total_hours') or 0
            fees_cents = row.get('total_fees_cents') or 0
            expenses_cents = row.get('total_expenses_cents') or 0
            total_billed_cents = fees_cents + expenses_cents
            fee_payments_cents = row.get('total_fee_payments_cents') or 0
            expense_payments_cents = row.get('total_expense_payments_cents') or 0
            total_payments_cents = fee_payments_cents + expense_payments_cents
            balance_cents = total_payments_cents - total_billed_cents

            values = [
                row.get('case_name') or '',
                row.get('client_name') or 'No Client',
                row.get('status') or 'Open',
                f"{total_hours:.1f}",
                f"${fees_cents / 100:.2f}",
                f"${expenses_cents / 100:.2f}",
                f"${total_billed_cents / 100:.2f}",
                f"${fee_payments_cents / 100:.2f}",
                f"${expense_payments_cents / 100:.2f}",
                f"${total_payments_cents / 100:.2f}",
                f"${balance_cents / 100:.2f}"
            ]

            for col_idx, value in enumerate(values):
                item = QTableWidgetItem(value)
                if col_idx >= 3:
                    item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
                if col_idx == 10:
                    if balance_cents < 0:
                        item.setForeground(QBrush(QColor("#dc3545")))
                    elif balance_cents > 0:
                        item.setForeground(QBrush(QColor("#28a745")))
                if row.get('status') == 'Closed':
                    item.setForeground(QBrush(QColor("#888888")))
                self.table.setItem(row_idx, col_idx, item)

        self.table.setSortingEnabled(True)

    def update_totals(self, totals: dict):
        self.matters_label.setText(f"Matters: {totals['matter_count']}")
        self.hours_label.setText(f"Hours: {totals['total_hours']:.1f}")
        self.billed_label.setText(f"Total Billed: ${totals['total_billed_cents'] / 100:.2f}")
        self.payments_label.setText(f"Total Payments: ${totals['total_payments_cents'] / 100:.2f}")

        balance_cents = totals['total_payments_cents'] - totals['total_billed_cents']
        balance_text = f"Net Balance: ${balance_cents / 100:.2f}"

        if balance_cents < 0:
            self.balance_label.setStyleSheet("font-weight: bold; color: #dc3545;")
        elif balance_cents > 0:
            self.balance_label.setStyleSheet("font-weight: bold; color: #28a745;")
        else:
            self.balance_label.setStyleSheet("font-weight: bold;")

        self.balance_label.setText(balance_text)

    def get_report_title(self) -> str:
        report_type = self.report_type_combo.currentData()
        if report_type == "monthly":
            year = self.year_combo.currentData()
            month = self.month_combo.currentData()
            month_name = date(year, month, 1).strftime("%B %Y")
            return f"Monthly Billing Summary - {month_name}"
        else:
            return "All-Time Billing Summary"

    def export_to_csv(self):
        if not self.current_data:
            QMessageBox.warning(self, "Warning", "Please generate a report first.")
            return

        report_title = self.get_report_title().replace(" ", "_").replace("-", "")
        default_filename = f"{report_title}.csv"

        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save CSV", default_filename, "CSV Files (*.csv)"
        )

        if not file_path:
            return

        try:
            import csv
            with open(file_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(self.HEADERS)

                for row_idx in range(self.table.rowCount()):
                    row_data = []
                    for col_idx in range(self.table.columnCount()):
                        item = self.table.item(row_idx, col_idx)
                        row_data.append(item.text() if item else "")
                    writer.writerow(row_data)

            QMessageBox.information(self, "Success", f"Report exported to:\n{file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to export CSV:\n{str(e)}")

    def export_to_docx(self):
        if not HAS_DOCX:
            QMessageBox.warning(self, "Missing Dependency", "python-docx is required for Word export.")
            return

        if not self.current_data:
            QMessageBox.warning(self, "Warning", "Please generate a report first.")
            return

        report_title = self.get_report_title()
        default_filename = f"{report_title.replace(' ', '_').replace('-', '')}.docx"

        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save Word Document", default_filename, "Word Documents (*.docx)"
        )

        if not file_path:
            return

        try:
            doc = Document()
            self._setup_document_style(doc)
            self._add_header(doc)

            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(report_title)
            title_run.bold = True
            title_run.font.size = Pt(16)

            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.add_run(f"Generated: {date.today().strftime('%B %d, %Y')}")

            doc.add_paragraph()

            table = doc.add_table(rows=1, cols=len(self.HEADERS))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            header_row = table.rows[0]
            for i, header in enumerate(self.HEADERS):
                cell = header_row.cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for row_idx in range(self.table.rowCount()):
                row_cells = table.add_row().cells
                for col_idx in range(self.table.columnCount()):
                    item = self.table.item(row_idx, col_idx)
                    text = item.text() if item else ""
                    row_cells[col_idx].text = text
                    if col_idx >= 3:
                        row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            doc.add_paragraph()

            totals = self._calculate_totals(self.current_data)
            summary_para = doc.add_paragraph()
            summary_para.add_run("Summary: ").bold = True
            summary_para.add_run(
                f"Matters: {totals['matter_count']} | "
                f"Hours: {totals['total_hours']:.1f} | "
                f"Total Billed: ${totals['total_billed_cents'] / 100:.2f} | "
                f"Total Payments: ${totals['total_payments_cents'] / 100:.2f} | "
                f"Net Balance: ${(totals['total_payments_cents'] - totals['total_billed_cents']) / 100:.2f}"
            )

            doc.save(file_path)
            QMessageBox.information(self, "Success", f"Report exported to:\n{file_path}")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to export Word document:\n{str(e)}")

    def _setup_document_style(self, doc):
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(9)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

    def _set_cell_border(self, cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def _add_header(self, doc):
        section = doc.sections[0]
        section.different_first_page_header_footer = True
        header = section.first_page_header
        header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        header_table.autofit = False
        left_cell = header_table.rows[0].cells[0]
        right_cell = header_table.rows[0].cells[1]
        self._set_cell_border(left_cell)
        self._set_cell_border(right_cell)
        left_cell.width = Inches(3.25)
        right_cell.width = Inches(3.25)

        if os.path.exists(self.image_path):
            left_para = left_cell.paragraphs[0]
            left_para.paragraph_format.space_before = Pt(0)
            left_para.paragraph_format.space_after = Pt(0)
            left_run = left_para.add_run()
            left_run.add_picture(self.image_path, width=Inches(2.5))

        right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        right_para.paragraph_format.space_before = Pt(0)
        right_para.paragraph_format.space_after = Pt(0)
        phone_run = right_para.add_run("404-556-7057")
        phone_run.font.size = Pt(10)
        right_para.add_run("\n")
        email_run = right_para.add_run("bbc@chintellalaw.com")
        email_run.font.size = Pt(10)
        right_para.add_run("\n")
        web_run = right_para.add_run("www.chintellalaw.com")
        web_run.font.size = Pt(10)

    def refresh(self):
        pass